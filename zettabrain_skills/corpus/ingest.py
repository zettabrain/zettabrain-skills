"""Document ingestion pipeline - reads, chunks, and prepares documents for the vector store."""

import hashlib
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from rich.console import Console

from zettabrain_skills.corpus.models import (
    CorpusDocument,
    CorpusManifest,
    DocumentChunk,
    DocumentStatus,
)

console = Console()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

METADATA_HEADER_PATTERN = re.compile(
    r"^=== CORPUS DOCUMENT ===\s*\n(.*?)\n={3,}\s*$",
    re.MULTILINE | re.DOTALL,
)


def read_pdf(file_path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def read_docx(file_path: Path) -> str:
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def read_text(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8")


READERS = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".txt": read_text,
    ".md": read_text,
}


def parse_metadata_header(content: str) -> Tuple[dict, str]:
    """Extract structured metadata header from document content.

    Returns (metadata_dict, remaining_content).
    """
    match = METADATA_HEADER_PATTERN.search(content)
    if not match:
        return {}, content

    header_block = match.group(1)
    metadata = {}
    for line in header_block.strip().split("\n"):
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.strip().lower().replace(" ", "_")
            value = value.strip()
            if value:
                metadata[key] = value

    body = content[match.end():].strip()
    return metadata, body


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> List[Tuple[str, int, int]]:
    """Split text into overlapping chunks.

    Returns list of (chunk_text, start_char, end_char).
    """
    if not text:
        return []

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = start + chunk_size

        if end < text_len:
            boundary = text.rfind("\n\n", start, end)
            if boundary == -1 or boundary <= start:
                boundary = text.rfind("\n", start, end)
            if boundary == -1 or boundary <= start:
                boundary = text.rfind(" ", start, end)
            if boundary > start:
                end = boundary

        actual_end = min(end, text_len)
        chunk = text[start:actual_end].strip()

        if chunk:
            chunks.append((chunk, start, actual_end))

        start = actual_end - chunk_overlap
        if start <= chunks[-1][1] if chunks else 0:
            start = actual_end

    return chunks


class DocumentIngestor:
    """Ingests documents into a corpus with incremental tracking."""

    def __init__(
        self,
        corpus_path: Path,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        self.corpus_path = Path(corpus_path)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.manifest_path = self.corpus_path / "manifest.json"
        self.manifest = self._load_manifest()

    def _load_manifest(self) -> CorpusManifest:
        if self.manifest_path.exists():
            data = json.loads(self.manifest_path.read_text())
            return CorpusManifest(**data)
        return CorpusManifest(corpus_id=str(uuid.uuid4()))

    def _save_manifest(self) -> None:
        self.corpus_path.mkdir(parents=True, exist_ok=True)
        self.manifest_path.write_text(
            self.manifest.model_dump_json(indent=2)
        )

    def ingest_file(self, file_path: Path) -> Optional[Tuple[CorpusDocument, List[DocumentChunk]]]:
        """Ingest a single file. Returns None if already ingested (unchanged)."""
        file_path = Path(file_path)

        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            console.print(f"[yellow]Skipping unsupported file: {file_path.name}[/yellow]")
            return None

        reader = READERS[file_path.suffix.lower()]
        raw_content = reader(file_path)

        content_hash = CorpusDocument.compute_hash(raw_content)
        if self.manifest.has_document(content_hash):
            console.print(f"[dim]Unchanged, skipping: {file_path.name}[/dim]")
            return None

        metadata, body = parse_metadata_header(raw_content)

        doc_id = str(uuid.uuid4())
        doc = CorpusDocument(
            id=doc_id,
            title=metadata.get("title", file_path.stem),
            source_path=str(file_path),
            content_hash=content_hash,
            issuing_body=metadata.get("issuing_body"),
            citation_ref=metadata.get("citation"),
            source_url=metadata.get("source_url"),
            effective_date=metadata.get("effective_date"),
            status=DocumentStatus(metadata.get("status", "current")),
            review_by=metadata.get("review_by"),
            corpus_version=metadata.get("corpus_version", "1"),
            file_type=file_path.suffix.lstrip("."),
            metadata=metadata,
        )

        text_to_chunk = body if body else raw_content
        raw_chunks = chunk_text(text_to_chunk, self.chunk_size, self.chunk_overlap)

        chunks = []
        for i, (chunk_text_content, start, end) in enumerate(raw_chunks):
            chunk = DocumentChunk(
                id=f"{doc_id}_{i}",
                document_id=doc_id,
                content=chunk_text_content,
                chunk_index=i,
                start_char=start,
                end_char=end,
                metadata={
                    "document_title": doc.title,
                    "source_path": str(file_path),
                    "chunk_index": str(i),
                },
            )
            chunks.append(chunk)

        doc.chunk_count = len(chunks)
        self.manifest.documents.append(doc)
        self.manifest.updated_at = datetime.now()
        self._save_manifest()

        console.print(
            f"[green]Ingested: {file_path.name} "
            f"({len(chunks)} chunks, {len(text_to_chunk)} chars)[/green]"
        )
        return doc, chunks

    def ingest_directory(self, directory: Path) -> List[Tuple[CorpusDocument, List[DocumentChunk]]]:
        """Ingest all supported files from a directory."""
        directory = Path(directory)
        if not directory.is_dir():
            raise FileNotFoundError(f"Directory not found: {directory}")

        results = []
        files = sorted(
            f for f in directory.rglob("*")
            if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
        )

        console.print(f"[blue]Found {len(files)} document(s) to process[/blue]")

        for file_path in files:
            result = self.ingest_file(file_path)
            if result:
                results.append(result)

        console.print(
            f"[green]Ingested {len(results)} new/updated document(s)[/green]"
        )
        return results

    def get_manifest(self) -> CorpusManifest:
        return self.manifest


