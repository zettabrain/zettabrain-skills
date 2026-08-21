"""
FastAPI web application for ZettaBrain Skills.
SPA architecture with JSON API + WebSocket streaming.
"""

import glob
import os
import time
import json
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from zettabrain_skills.core.engine import GenerationEngine
from zettabrain_skills.core.models import GenerationRequest
from zettabrain_skills.skills.parser import load_skill
from zettabrain_skills.web.document_store import DocumentStore

app = FastAPI(
    title="ZettaBrain Skills",
    description="AI-powered document generation with skills",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).parent
STATIC_DIR = BASE_DIR / "static"

if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

SKILLS_DIR = os.getenv("SKILLS_DIR", "examples")
CORPUS_PATH = os.getenv("CORPUS_PATH", ".corpus")
SETTINGS_FILE = os.getenv("SETTINGS_FILE", ".data/settings.json")

document_store = DocumentStore()


def _load_persisted_settings():
    """Load settings from disk and apply to environment on startup."""
    settings_path = Path(SETTINGS_FILE)
    if not settings_path.exists():
        return
    try:
        data = json.loads(settings_path.read_text())
        env_map = {
            "llm_provider": "LLM_PROVIDER",
            "ollama_url": "OLLAMA_BASE_URL",
            "ollama_model": "OLLAMA_MODEL",
            "ollama_timeout": "OLLAMA_TIMEOUT",
            "groq_api_key": "GROQ_API_KEY",
            "groq_model": "GROQ_MODEL",
            "together_api_key": "TOGETHER_API_KEY",
            "together_model": "TOGETHER_MODEL",
            "aws_access_key": "AWS_ACCESS_KEY_ID",
            "aws_secret_key": "AWS_SECRET_ACCESS_KEY",
            "aws_region": "AWS_REGION",
            "bedrock_model": "BEDROCK_MODEL",
            "embedding_model": "EMBEDDING_MODEL",
            "embedding_url": "OLLAMA_HOST",
        }
        for key, env_var in env_map.items():
            val = data.get(key)
            if val and not os.getenv(env_var):
                os.environ[env_var] = str(val)
    except Exception:
        pass


def _save_persisted_settings(settings: Dict[str, Any]):
    """Save current settings to disk for persistence across reinstalls."""
    settings_path = Path(SETTINGS_FILE)
    settings_path.parent.mkdir(parents=True, exist_ok=True)
    existing = {}
    if settings_path.exists():
        try:
            existing = json.loads(settings_path.read_text())
        except Exception:
            pass
    existing.update({k: v for k, v in settings.items() if v})
    settings_path.write_text(json.dumps(existing, indent=2))


_load_persisted_settings()


def _load_skills() -> List[Dict[str, Any]]:
    import re
    skills = []
    for skill_file in sorted(glob.glob(f"{SKILLS_DIR}/*.md")):
        try:
            skill = load_skill(skill_file)
            # Detect {{variables}} in instructions
            variables = list(set(re.findall(r"\{\{(\w+)\}\}", skill.instructions)))
            skills.append({
                "file": skill_file,
                "name": skill.name,
                "display_name": skill.name.replace("-", " ").title(),
                "description": skill.description,
                "business_type": skill.business_type,
                "version": skill.version,
                "requires_corpus": skill.requires_corpus,
                "citation_required": skill.citation_required,
                "variables": variables,
            })
        except Exception:
            continue
    return skills


def _get_engine() -> GenerationEngine:
    engine = GenerationEngine()
    corpus_path = Path(CORPUS_PATH)
    if corpus_path.exists():
        try:
            from zettabrain_skills.corpus.retrieval import CorpusRetriever
            engine.corpus_retriever = CorpusRetriever(corpus_path=str(corpus_path))
        except Exception:
            pass
    return engine


# ── Root: serve SPA ──────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def root():
    index = STATIC_DIR / "index.html"
    if index.exists():
        return HTMLResponse(index.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>ZettaBrain Skills</h1><p>Static files not found.</p>")


# ── API: Status ──────────────────────────────────────────

@app.get("/api/status")
async def api_status():
    from zettabrain_skills.llm.factory import create_llm_provider, get_provider_info

    provider = create_llm_provider()
    llm_healthy = provider.check_health()
    model_info = provider.get_model_info() if llm_healthy else {}
    provider_info = get_provider_info()

    corpus_chunks = 0
    corpus_docs = 0
    try:
        from zettabrain_skills.corpus.retrieval import CorpusRetriever
        corpus_path = Path(CORPUS_PATH)
        if corpus_path.exists() and (corpus_path / "manifest.json").exists():
            retriever = CorpusRetriever(corpus_path=str(corpus_path))
            corpus_chunks = retriever.chunk_count
            corpus_docs = retriever.document_count
    except Exception:
        pass

    return {
        "llm": {
            "healthy": llm_healthy,
            "provider": provider_info.get("provider", "unknown"),
            "model": model_info.get("model", "unknown"),
        },
        "skills": {
            "count": len(_load_skills()),
            "directory": SKILLS_DIR,
        },
        "corpus": {
            "chunks": corpus_chunks,
            "documents": corpus_docs,
            "path": CORPUS_PATH,
        },
        "documents_generated": document_store.count(),
        "version": "1.0.0",
    }


# ── API: Skills ──────────────────────────────────────────

@app.get("/api/skills")
async def api_skills():
    return _load_skills()


# ── API: Generate ────────────────────────────────────────

class GenerateRequest(BaseModel):
    skill_file: str
    input_text: str
    customer_name: str = ""
    customer_email: str = ""
    customer_phone: str = ""


@app.post("/api/generate")
async def api_generate(req: GenerateRequest):
    skill_path = Path(req.skill_file)
    if not skill_path.exists():
        raise HTTPException(status_code=400, detail=f"Skill file not found: {req.skill_file}")

    skill = load_skill(str(skill_path))
    engine = _get_engine()

    if not engine.llm_provider.check_health():
        raise HTTPException(status_code=503, detail="LLM provider is not running")

    today = datetime.now()
    parts = [f"TODAY'S DATE: {today.strftime('%B %d, %Y')}\n\n"]
    if req.customer_name:
        parts.append(f"Customer/Client: {req.customer_name}\n")
    if req.customer_email:
        parts.append(f"Email: {req.customer_email}\n")
    if req.customer_phone:
        parts.append(f"Phone: {req.customer_phone}\n")
    if req.customer_name or req.customer_email or req.customer_phone:
        parts.append("\n")
    parts.append(req.input_text)

    gen_request = GenerationRequest(
        input="".join(parts),
        skill_name=skill.name,
        business_id="default",
    )

    result = engine.generate(skill, gen_request)

    if not result.success:
        raise HTTPException(status_code=500, detail=f"Generation failed: {result.error}")

    doc_data = {
        "id": result.id,
        "skill_name": skill.name,
        "skill_display": skill.name.replace("-", " ").title(),
        "customer_name": req.customer_name or "Customer",
        "customer_email": req.customer_email,
        "customer_phone": req.customer_phone,
        "request": req.input_text,
        "content": result.content,
        "citations": result.citations,
        "created_at": result.created_at.strftime("%Y-%m-%d %H:%M:%S"),
        "generation_time_ms": result.generation_time_ms,
    }
    document_store.insert(doc_data)

    return doc_data


# ── WebSocket: Streaming Generation ─────────────────────

@app.websocket("/ws/generate")
async def ws_generate(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            data = await websocket.receive_json()

            skill_file = data.get("skill_file")
            input_text = data.get("input_text", "")
            customer_name = data.get("customer_name", "")

            if not skill_file or not input_text:
                await websocket.send_json({"type": "error", "message": "skill_file and input_text required"})
                continue

            skill_path = Path(skill_file)
            if not skill_path.exists():
                await websocket.send_json({"type": "error", "message": f"Skill not found: {skill_file}"})
                continue

            try:
                skill = load_skill(str(skill_path))
                engine = _get_engine()

                if not engine.llm_provider.check_health():
                    await websocket.send_json({"type": "error", "message": "LLM provider is not running"})
                    continue

                await websocket.send_json({"type": "status", "message": "Generating..."})

                today = datetime.now()
                parts = [f"TODAY'S DATE: {today.strftime('%B %d, %Y')}\n\n"]
                if customer_name:
                    parts.append(f"Customer/Client: {customer_name}\n\n")
                parts.append(input_text)

                gen_request = GenerationRequest(
                    input="".join(parts),
                    skill_name=skill.name,
                    business_id="default",
                )

                # Stream tokens if provider supports it
                start_time = time.time()
                corpus_context = None
                citations = []

                if skill.requires_corpus and engine.corpus_retriever:
                    corpus_text, citation_objects = (
                        engine.corpus_retriever.get_context_for_generation(
                            query=input_text, n_results=5, min_relevance=0.3,
                            business_type=skill.business_type if skill.business_type != "generic" else None,
                        )
                    )
                    if corpus_text:
                        corpus_context = corpus_text
                        citations = [
                            f"{c.document_title}"
                            + (f" ({c.citation_ref})" if c.citation_ref else "")
                            for c in citation_objects
                        ]
                        await websocket.send_json({"type": "citations", "citations": citations})

                prompt = engine.build_prompt(skill, "".join(parts), gen_request.context, corpus_context)

                temperature = skill.temperature
                max_tokens = skill.max_tokens

                full_content = ""
                try:
                    for token in engine.llm_provider.stream(
                        prompt=prompt, temperature=temperature, max_tokens=max_tokens
                    ):
                        full_content += token
                        await websocket.send_json({"type": "token", "token": token})
                except NotImplementedError:
                    full_content = engine.llm_provider.generate(
                        prompt=prompt, temperature=temperature, max_tokens=max_tokens
                    )
                    await websocket.send_json({"type": "token", "token": full_content})

                generation_time_ms = int((time.time() - start_time) * 1000)

                import uuid
                doc_id = str(uuid.uuid4())
                doc_data = {
                    "id": doc_id,
                    "skill_name": skill.name,
                    "skill_display": skill.name.replace("-", " ").title(),
                    "customer_name": customer_name or "Customer",
                    "request": input_text,
                    "content": full_content,
                    "citations": citations,
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "generation_time_ms": generation_time_ms,
                }
                document_store.insert(doc_data)

                await websocket.send_json({
                    "type": "done",
                    "id": doc_id,
                    "skill": skill.name,
                    "generation_time_ms": generation_time_ms,
                    "model": getattr(engine.llm_provider, "model", "unknown"),
                    "citations": citations,
                })

            except Exception as e:
                await websocket.send_json({"type": "error", "message": str(e)})

    except WebSocketDisconnect:
        pass


# ── API: Documents ───────────────────────────────────────

@app.get("/api/documents")
async def api_documents():
    return document_store.get_all(limit=50)


@app.get("/api/documents/{doc_id}")
async def api_document(doc_id: str):
    doc = document_store.get_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@app.delete("/api/documents/{doc_id}")
async def api_delete_document(doc_id: str):
    doc = document_store.get_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    document_store.delete(doc_id)
    return {"deleted": True, "id": doc_id}


class RefineRequest(BaseModel):
    instruction: str
    doc_id: str


@app.post("/api/documents/refine")
async def api_refine_document(req: RefineRequest):
    """Refine/modify an existing generated document with a follow-up instruction."""
    from zettabrain_skills.llm.factory import create_llm_provider

    doc = document_store.get_by_id(req.doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    provider = create_llm_provider()
    if not provider.check_health():
        raise HTTPException(status_code=503, detail="LLM provider is not running")

    refine_prompt = f"""You are a document editor. Your ONLY job is to revise the document below based on the user's instruction. Do NOT refuse. Do NOT ask clarifying questions. Just apply the requested change and output the revised document.

ORIGINAL DOCUMENT:
{doc['content']}

USER'S EDIT INSTRUCTION: {req.instruction}

RULES:
- Apply the change directly. If the user says "1500 sq ft" they mean change the square footage to 1500.
- Keep all other content the same unless the change logically requires adjustments (e.g., changing area may affect cost calculations).
- Output the COMPLETE revised document with the change applied.
- Do NOT add commentary, do NOT explain what you changed. Just output the document.

REVISED DOCUMENT:"""

    start_time = time.time()
    try:
        content = provider.generate(prompt=refine_prompt, temperature=0.3, max_tokens=4000)
        generation_time_ms = int((time.time() - start_time) * 1000)

        import uuid
        new_id = str(uuid.uuid4())
        doc_data = {
            "id": new_id,
            "skill_name": doc.get("skill_name", "refined"),
            "skill_display": doc.get("skill_display", "Refined Document"),
            "customer_name": doc.get("customer_name", "Customer"),
            "customer_email": doc.get("customer_email", ""),
            "request": f"Refined: {req.instruction}",
            "content": content,
            "citations": doc.get("citations", []),
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "generation_time_ms": generation_time_ms,
        }
        document_store.insert(doc_data)

        return {
            "id": new_id,
            "content": content,
            "generation_time_ms": generation_time_ms,
            "model": getattr(provider, "model", "unknown"),
            "original_id": req.doc_id,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Refinement failed: {e}")


@app.get("/api/documents/{doc_id}/pdf")
async def api_document_pdf(doc_id: str):
    from fpdf import FPDF

    doc = document_store.get_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    def sanitize(text: str) -> str:
        replacements = {
            "•": "-", "’": "'", "‘": "'",
            "“": '"', "”": '"', "–": "-",
            "—": "--", "…": "...", " ": " ",
            "‐": "-", "‑": "-", "‒": "-",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    customer_name = doc.get("customer_name", "")

    # Letterhead: Customer/Company name
    if customer_name and customer_name != "Customer":
        pdf.set_font("Helvetica", "B", 22)
        pdf.set_text_color(33, 37, 41)
        pdf.cell(0, 12, sanitize(customer_name), ln=True, align="C")
        pdf.ln(2)

    # Document type subtitle
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(99, 102, 241)
    pdf.cell(0, 8, sanitize(doc.get("skill_display", "Document")), ln=True, align="C")
    pdf.ln(3)
    pdf.set_draw_color(99, 102, 241)
    pdf.set_line_width(0.6)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    # Metadata row
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(120, 120, 120)
    meta_parts = []
    if doc.get("created_at"):
        meta_parts.append(f"Date: {doc['created_at']}")
    meta_parts.append(f"Ref: {doc['id'][:12]}")
    pdf.cell(0, 5, " | ".join(meta_parts), ln=True, align="C")
    pdf.ln(8)

    # Content
    left_margin = pdf.l_margin
    pdf.set_text_color(33, 33, 33)
    for line in doc["content"].split("\n"):
        if line.startswith("### "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_x(left_margin)
            pdf.multi_cell(0, 6, sanitize(line[4:]))
            pdf.ln(1)
        elif line.startswith("## "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_x(left_margin)
            pdf.multi_cell(0, 7, sanitize(line[3:]))
            pdf.ln(1)
        elif line.startswith("# "):
            pdf.ln(5)
            pdf.set_font("Helvetica", "B", 15)
            pdf.set_x(left_margin)
            pdf.multi_cell(0, 8, sanitize(line[2:]))
            pdf.ln(2)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_x(left_margin + 6)
            pdf.multi_cell(0, 5, "- " + sanitize(line[2:]))
        elif line.strip().startswith("|") and "|" in line[1:]:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_x(left_margin)
            pdf.multi_cell(0, 5, sanitize(line.strip()))
        elif line.strip() == "":
            pdf.ln(3)
        else:
            clean = line.replace("**", "").replace("*", "").strip()
            if clean:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_x(left_margin)
                pdf.multi_cell(0, 5, sanitize(clean))

    # Citations
    if doc.get("citations"):
        pdf.ln(8)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(99, 102, 241)
        pdf.set_x(left_margin)
        pdf.cell(0, 7, "Sources", ln=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(80, 80, 80)
        for citation in doc["citations"]:
            pdf.set_x(left_margin + 4)
            pdf.multi_cell(0, 5, "- " + sanitize(citation))

    # Footer
    pdf.ln(10)
    pdf.set_draw_color(220, 220, 220)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(170, 170, 170)
    pdf.cell(0, 4, "Generated with ZettaBrain Skills | zettabrain.ai", ln=True, align="C")

    pdf_output = pdf.output()
    skill_name = doc.get("skill_name", "document").replace(" ", "-")
    return Response(
        content=bytes(pdf_output),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=zettabrain-{skill_name}-{doc_id[:8]}.pdf"},
    )


@app.get("/api/documents/{doc_id}/docx")
async def api_document_docx(doc_id: str):
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = document_store.get_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    docx_doc = DocxDocument()

    style = docx_doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    customer_name = doc.get("customer_name", "")

    # Letterhead: Customer/Company name
    if customer_name and customer_name != "Customer":
        heading = docx_doc.add_heading(customer_name, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x21, 0x25, 0x29)

    # Document type subtitle
    subtitle = docx_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(doc.get("skill_display", "Document"))
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    run.font.size = Pt(12)

    # Metadata line
    meta_para = docx_doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"Date: {doc.get('created_at', '')} | Ref: {doc['id'][:12]}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    docx_doc.add_paragraph()

    for line in doc["content"].split("\n"):
        if line.startswith("# "):
            docx_doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            docx_doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            docx_doc.add_heading(line[4:], level=3)
        elif line.strip():
            docx_doc.add_paragraph(line)

    if doc.get("citations"):
        docx_doc.add_paragraph()
        docx_doc.add_heading("Sources", level=2)
        for citation in doc["citations"]:
            docx_doc.add_paragraph(citation, style="List Bullet")

    footer = docx_doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Generated with ZettaBrain Skills | zettabrain.ai")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    docx_bytes = BytesIO()
    docx_doc.save(docx_bytes)
    docx_bytes.seek(0)

    skill_name = doc.get("skill_name", "document").replace(" ", "-")
    return Response(
        content=docx_bytes.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=zettabrain-{skill_name}-{doc_id[:8]}.docx"},
    )


# ── API: Corpus ──────────────────────────────────────────

@app.get("/api/corpus/status")
async def api_corpus_status():
    try:
        from zettabrain_skills.corpus.retrieval import CorpusRetriever
        corpus_path = Path(CORPUS_PATH)
        if not corpus_path.exists() or not (corpus_path / "manifest.json").exists():
            return {"configured": False, "chunks": 0, "documents": 0}

        retriever = CorpusRetriever(corpus_path=str(corpus_path))
        manifest = retriever.ingestor.get_manifest()

        docs = [
            {"title": d.title, "type": d.file_type, "chunks": d.chunk_count, "status": d.status.value}
            for d in manifest.documents
        ]
        return {
            "configured": True,
            "chunks": retriever.chunk_count,
            "documents": retriever.document_count,
            "document_list": docs,
        }
    except Exception as e:
        return {"configured": False, "error": str(e)}


class CorpusSearchRequest(BaseModel):
    query: str
    n_results: int = 5


@app.post("/api/corpus/search")
async def api_corpus_search(req: CorpusSearchRequest):
    from zettabrain_skills.corpus.retrieval import CorpusRetriever

    corpus_path = Path(CORPUS_PATH)
    if not corpus_path.exists():
        raise HTTPException(status_code=400, detail="No corpus configured")

    retriever = CorpusRetriever(corpus_path=str(corpus_path))
    results = retriever.search(req.query, n_results=req.n_results)

    return [
        {
            "title": r.citation.document_title,
            "excerpt": r.chunk.content[:300],
            "score": round(r.score, 3),
            "citation_ref": r.citation.citation_ref,
            "issuing_body": r.citation.issuing_body,
        }
        for r in results
    ]


# ── API: Corpus Upload & Ingest ─────────────────────────

CORPUS_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@app.post("/api/corpus/upload")
async def api_corpus_upload(files: List[UploadFile] = File(...)):
    """Upload documents to the corpus uploads directory (does not ingest)."""
    corpus_path = Path(CORPUS_PATH)
    upload_dir = corpus_path / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)

    saved_files = []
    skipped_files = []

    for file in files:
        ext = Path(file.filename).suffix.lower()
        if ext not in CORPUS_SUPPORTED_EXTENSIONS:
            skipped_files.append({"name": file.filename, "reason": f"Unsupported format: {ext}"})
            continue

        dest = upload_dir / file.filename
        content = await file.read()
        dest.write_bytes(content)
        saved_files.append(file.filename)

    return {
        "uploaded": len(saved_files),
        "files": saved_files,
        "skipped": skipped_files,
        "message": f"Uploaded {len(saved_files)} file(s). Click 'Ingest' to process them.",
    }


@app.post("/api/corpus/ingest")
async def api_corpus_ingest():
    """Ingest all uploaded corpus documents into the vector store."""
    from zettabrain_skills.corpus.retrieval import CorpusRetriever

    corpus_path = Path(CORPUS_PATH)
    upload_dir = corpus_path / "uploads"

    if not upload_dir.exists() or not any(upload_dir.iterdir()):
        raise HTTPException(status_code=400, detail="No files to ingest. Upload documents first.")

    embedding_model = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
    retriever = CorpusRetriever(corpus_path=str(corpus_path), embedding_model=embedding_model)
    ingested_count = 0
    errors = []

    for file_path in sorted(upload_dir.rglob("*")):
        if file_path.is_file() and file_path.suffix.lower() in CORPUS_SUPPORTED_EXTENSIONS:
            try:
                count = retriever.ingest(file_path)
                ingested_count += count
            except Exception as e:
                errors.append({"file": file_path.name, "error": str(e)})

    return {
        "ingested": ingested_count,
        "total_documents": retriever.document_count,
        "total_chunks": retriever.chunk_count,
        "errors": errors,
    }


# ── API: Skills Upload ──────────────────────────────────

@app.post("/api/skills/upload")
async def api_skills_upload(file: UploadFile = File(...)):
    """Upload a new skill file (.md with YAML frontmatter)."""
    if not file.filename.endswith(".md"):
        raise HTTPException(status_code=400, detail="Skill files must be .md format")

    content = await file.read()
    text = content.decode("utf-8")

    if not text.startswith("---"):
        raise HTTPException(
            status_code=400,
            detail="Invalid skill file: must start with YAML frontmatter (---)",
        )

    skills_dir = Path(SKILLS_DIR)
    skills_dir.mkdir(parents=True, exist_ok=True)

    dest = skills_dir / file.filename
    dest.write_text(text, encoding="utf-8")

    try:
        skill = load_skill(str(dest))
        return {
            "success": True,
            "file": str(dest),
            "name": skill.name,
            "display_name": skill.name.replace("-", " ").title(),
            "description": skill.description,
            "version": skill.version,
            "requires_corpus": skill.requires_corpus,
        }
    except Exception as e:
        dest.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Invalid skill file: {e}")


@app.delete("/api/skills/{skill_name}")
async def api_delete_skill(skill_name: str):
    """Delete a skill file by name."""
    skills_dir = Path(SKILLS_DIR)
    for skill_file in skills_dir.glob("*.md"):
        try:
            skill = load_skill(str(skill_file))
            if skill.name == skill_name:
                skill_file.unlink()
                return {"deleted": True, "name": skill_name}
        except Exception:
            continue
    raise HTTPException(status_code=404, detail=f"Skill not found: {skill_name}")


class SkillBuilderRequest(BaseModel):
    business_name: str
    document_type: str
    description: str
    sections: str = ""
    rules: str = ""
    uses_corpus: bool = False
    tone: str = "professional"


@app.post("/api/skills/build")
async def api_skills_build(req: SkillBuilderRequest):
    """Use the LLM to generate a skill file from a plain-language description."""
    from zettabrain_skills.llm.factory import create_llm_provider

    provider = create_llm_provider()
    if not provider.check_health():
        raise HTTPException(status_code=503, detail="LLM provider is not running")

    skill_prompt = f"""You are a skill file generator for ZettaBrain Skills, an AI document generation platform.

Generate a complete skill file in markdown format with YAML frontmatter based on the user's description.

The skill file format is:
```
---
name: skill-name-in-kebab-case
version: 1.0.0
description: One-line description of what the skill generates
business_type: the industry/type (e.g., construction, legal, medical, finance, retail, generic)
requires_corpus: true/false
citation_required: true/false
temperature: 0.3 to 0.9 (lower = more precise, higher = more creative)
max_tokens: 2000 to 8000
tags:
  - relevant
  - tags
---

# Skill Title

You are a [role] generating [document type] for [business context].

## Your Task

[Clear description of what to generate]

## Required Sections

[Numbered list of sections the document must include]

## Formatting Rules

[Rules about format, style, and presentation]

## Important Constraints

[Business rules, compliance requirements, things to always/never do]
```

USER'S REQUEST:
- Business: {req.business_name}
- Document type: {req.document_type}
- Description: {req.description}
- Sections they want: {req.sections or 'Let the AI decide appropriate sections'}
- Business rules: {req.rules or 'None specified'}
- Uses reference corpus: {req.uses_corpus}
- Tone: {req.tone}

Generate ONLY the skill file content (starting with ---). Make it comprehensive and professional.
Do not wrap it in a code block. Output the raw skill file content directly."""

    try:
        content = provider.generate(prompt=skill_prompt, temperature=0.4, max_tokens=4000)

        # Clean up any code fences the LLM might add
        content = content.strip()
        if content.startswith("```"):
            content = "\n".join(content.split("\n")[1:])
        if content.endswith("```"):
            content = "\n".join(content.split("\n")[:-1])
        content = content.strip()

        if not content.startswith("---"):
            raise HTTPException(status_code=500, detail="LLM did not generate a valid skill file")

        # Generate a filename from the skill name
        import re
        name_match = re.search(r"^name:\s*(.+)$", content, re.MULTILINE)
        filename = "custom-skill.md"
        if name_match:
            filename = name_match.group(1).strip() + ".md"

        return {
            "content": content,
            "filename": filename,
            "preview": True,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Skill generation failed: {e}")


@app.post("/api/skills/save")
async def api_skills_save(data: Dict[str, Any]):
    """Save a generated skill file to disk."""
    content = data.get("content", "")
    filename = data.get("filename", "custom-skill.md")

    if not content.startswith("---"):
        raise HTTPException(status_code=400, detail="Invalid skill content")

    if not filename.endswith(".md"):
        filename += ".md"

    skills_dir = Path(SKILLS_DIR)
    skills_dir.mkdir(parents=True, exist_ok=True)

    dest = skills_dir / filename
    dest.write_text(content, encoding="utf-8")

    try:
        skill = load_skill(str(dest))
        return {
            "success": True,
            "file": str(dest),
            "name": skill.name,
            "display_name": skill.name.replace("-", " ").title(),
            "description": skill.description,
        }
    except Exception as e:
        dest.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Generated skill has errors: {e}")


# ── API: Settings ───────────────────────────────────────

@app.get("/api/settings")
async def api_get_settings():
    """Get current LLM and embedding configuration for all providers."""
    return {
        "llm_provider": os.getenv("LLM_PROVIDER", "ollama"),
        # Ollama
        "ollama_url": os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
        "ollama_model": os.getenv("OLLAMA_MODEL", "llama3.1:8b"),
        "ollama_timeout": os.getenv("OLLAMA_TIMEOUT", "600"),
        # Groq
        "groq_api_key": bool(os.getenv("GROQ_API_KEY")),
        "groq_model": os.getenv("GROQ_MODEL", "llama-3.1-8b-instant"),
        # Together
        "together_api_key": bool(os.getenv("TOGETHER_API_KEY")),
        "together_model": os.getenv("TOGETHER_MODEL", "meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo"),
        # Bedrock
        "aws_access_key": bool(os.getenv("AWS_ACCESS_KEY_ID")),
        "aws_region": os.getenv("AWS_REGION", "us-east-1"),
        "bedrock_model": os.getenv("BEDROCK_MODEL", "meta.llama3-1-8b-instruct-v1:0"),
        # Embedding
        "embedding_model": os.getenv("EMBEDDING_MODEL", "nomic-embed-text"),
        "embedding_url": os.getenv("OLLAMA_HOST", ""),
        # Status
        "configured_providers": {
            "ollama": True,
            "groq": bool(os.getenv("GROQ_API_KEY")),
            "together": bool(os.getenv("TOGETHER_API_KEY")),
            "bedrock": bool(os.getenv("AWS_ACCESS_KEY_ID")),
        },
    }


@app.post("/api/settings")
async def api_update_settings(settings: Dict[str, Any]):
    """Update LLM and embedding configuration (runtime, not persisted across restarts)."""
    updated = []

    provider = settings.get("llm_provider")
    if provider:
        os.environ["LLM_PROVIDER"] = provider
        updated.append(f"provider={provider}")

    # Ollama settings
    if settings.get("ollama_url"):
        os.environ["OLLAMA_BASE_URL"] = settings["ollama_url"]
        updated.append("ollama_url")
    if settings.get("ollama_model"):
        os.environ["OLLAMA_MODEL"] = settings["ollama_model"]
        updated.append(f"model={settings['ollama_model']}")
    if settings.get("ollama_timeout"):
        os.environ["OLLAMA_TIMEOUT"] = str(settings["ollama_timeout"])
        updated.append("timeout")

    # Groq settings
    if settings.get("groq_api_key"):
        os.environ["GROQ_API_KEY"] = settings["groq_api_key"]
        updated.append("groq_api_key")
    if settings.get("groq_model"):
        os.environ["GROQ_MODEL"] = settings["groq_model"]
        os.environ["OLLAMA_MODEL"] = settings["groq_model"]
        updated.append(f"model={settings['groq_model']}")

    # Together settings
    if settings.get("together_api_key"):
        os.environ["TOGETHER_API_KEY"] = settings["together_api_key"]
        updated.append("together_api_key")
    if settings.get("together_model"):
        os.environ["TOGETHER_MODEL"] = settings["together_model"]
        updated.append(f"model={settings['together_model']}")

    # Bedrock settings
    if settings.get("aws_access_key"):
        os.environ["AWS_ACCESS_KEY_ID"] = settings["aws_access_key"]
        updated.append("aws_access_key")
    if settings.get("aws_secret_key"):
        os.environ["AWS_SECRET_ACCESS_KEY"] = settings["aws_secret_key"]
        updated.append("aws_secret_key")
    if settings.get("aws_region"):
        os.environ["AWS_REGION"] = settings["aws_region"]
        updated.append(f"region={settings['aws_region']}")
    if settings.get("bedrock_model"):
        os.environ["BEDROCK_MODEL"] = settings["bedrock_model"]
        updated.append(f"model={settings['bedrock_model']}")

    # Embedding settings
    if settings.get("embedding_model"):
        os.environ["EMBEDDING_MODEL"] = settings["embedding_model"]
        updated.append(f"embedding={settings['embedding_model']}")
    if settings.get("embedding_url"):
        os.environ["OLLAMA_HOST"] = settings["embedding_url"]
        updated.append("embedding_url")

    # Persist to disk so settings survive reinstalls
    _save_persisted_settings(settings)

    return {
        "updated": updated,
        "message": f"Settings saved and persisted: {', '.join(updated) if updated else 'no changes'}",
    }


# ── Health ───────────────────────────────────────────────

@app.get("/health")
async def health_check():
    from zettabrain_skills.llm.factory import create_llm_provider
    provider = create_llm_provider()
    healthy = provider.check_health()
    return {"status": "healthy" if healthy else "degraded", "version": "1.0.0"}
